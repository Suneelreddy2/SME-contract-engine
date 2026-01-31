"""Extract text from PDF, DOCX, and plain text files for contract analysis."""
from typing import Tuple
import io


def extract_from_pdf(content: bytes) -> Tuple[str, str | None]:
    """Extract text from PDF bytes. Returns (text, error_message)."""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=content, filetype="pdf")
        parts = []
        for page in doc:
            parts.append(page.get_text())
        doc.close()
        text = "\n".join(parts).strip()
        if not text:
            return "", "PDF appears to be image-based or empty; no text could be extracted."
        return text, None
    except Exception as e:
        return "", str(e)


def extract_from_docx(content: bytes) -> Tuple[str, str | None]:
    """Extract text from DOCX bytes. Returns (text, error_message)."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" ".join(cell.text for cell in row.cells))
        text = "\n".join(parts).strip()
        if not text:
            return "", "DOCX appears to be empty."
        return text, None
    except Exception as e:
        return "", str(e)


def extract_from_txt(content: bytes) -> Tuple[str, str | None]:
    """Decode plain text. Returns (text, error_message)."""
    try:
        for enc in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                return content.decode(enc).strip(), None
            except UnicodeDecodeError:
                continue
        return "", "Could not decode file as text."
    except Exception as e:
        return "", str(e)


def extract_text_from_file(content: bytes, filename: str) -> Tuple[str, str | None]:
    """Route by extension. Returns (extracted_text, error_message)."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return extract_from_pdf(content)
    if name.endswith(".docx"):
        return extract_from_docx(content)
    if name.endswith(".txt"):
        return extract_from_txt(content)
    return "", f"Unsupported format. Use PDF, DOCX, or TXT. Got: {filename}"
